#!/usr/bin/env python3
"""
generate-avalere-docx build engine (template-as-base).

Starts from the pinned Avalere Word template, clears its demo body content while
keeping the section properties (so headers, footers, and the embedded logo are
preserved), and adds paragraphs using the template's own named styles. The brand
comes from the template's styles and theme, never from ad-hoc formatting.

Usage:
    python3 build_doc.py <spec.json> [--template <path>] [--out <path>]
    python3 build_doc.py list-styles [<template>]

Spec JSON shape:
{
  "output": "doc.docx",
  "blocks": [
    {"style": "Title",     "text": "Velorixa"},
    {"style": "Subtitle",  "text": "Brand strategy kickoff recap"},
    {"style": "Heading 1", "text": "Brand ambition"},
    {"style": "Normal",    "text": "The core pressure is clarity.", "citations": ["623851f2"]},
    {"style": "List Bullet","text": "Four phases ...", "citations": ["623851f2"]}
  ]
}
"""
import sys, os, json, argparse, re

from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.normpath(os.path.join(HERE, "..", "_assets", "Avalere_Doc_template.docx"))


def _norm(s):
    s = s.replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", s).strip().lower()


def style_index(doc):
    """Map normalised style name and style_id -> actual style name python-docx accepts."""
    idx = {}
    for st in doc.styles:
        try:
            nm = st.name
        except Exception:
            nm = None
        if nm:
            idx[_norm(nm)] = nm
        if getattr(st, "style_id", None):
            idx.setdefault(_norm(st.style_id), nm or st.style_id)
    return idx


def resolve_style(idx, requested, default="Normal"):
    r = _norm(requested)
    if r in idx:
        return idx[r]
    for k, v in idx.items():
        if r in k:
            return v
    return default


def clear_body(doc):
    """Remove existing body content but keep the trailing sectPr (headers/footers/logo, page setup)."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def cite_suffix(citations):
    return (" " + " ".join(f"[src:{c}]" for c in citations)) if citations else ""


def build(spec, template_path, out_override=None):
    doc = Document(template_path)
    idx = style_index(doc)
    if spec.get("clear_template_body", True):
        clear_body(doc)
    for block in spec.get("blocks", []):
        style = resolve_style(idx, block.get("style", "Normal"))
        text = block.get("text", "") + cite_suffix(block.get("citations", []))
        doc.add_paragraph(text, style=style)
    out = out_override or spec.get("output", "document.docx")
    if not os.path.isabs(out):
        out = os.path.join(os.getcwd(), out)
    doc.save(out)
    return out, len(spec.get("blocks", []))


def main():
    if len(sys.argv) >= 2 and sys.argv[1] == "list-styles":
        tpl = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_TEMPLATE
        doc = Document(tpl)
        for st in doc.styles:
            try:
                print(f"{st.type}\t{st.style_id}\t{st.name}")
            except Exception:
                pass
        return
    ap = argparse.ArgumentParser()
    ap.add_argument("spec")
    ap.add_argument("--template", default=DEFAULT_TEMPLATE)
    ap.add_argument("--out", default=None)
    args = ap.parse_args()
    spec = json.load(open(args.spec))
    out, n = build(spec, args.template, args.out)
    print(f"Wrote {out} ({n} blocks) from {os.path.basename(args.template)}")


if __name__ == "__main__":
    main()
